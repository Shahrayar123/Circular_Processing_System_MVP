"""The Word working document — one per circular.

The audit team's own convention, taken from the annotated circulars ABL supplied: a
short note against each actionable clause naming the proposal it produced, and
"Information" against the rest.

The real system anchors these as true Word comments through an OOXML helper. The MVP
puts them in a right-hand column of a two-column table, which renders the same idea
without the OOXML work.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from . import config, store

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x2E, 0x7A, 0x4F)
AMBER = RGBColor(0x9C, 0x6F, 0x11)
RED = RGBColor(0xB0, 0x3A, 0x30)

COLOURS = {"New": GREEN, "Amendment": AMBER, "Deletion": RED, "No action": GREY}


def _note(proposal: dict | None) -> tuple[str, RGBColor]:
    if not proposal:
        return "Information", GREY
    change = proposal["change_type"]
    if change == "New":
        return f"New test {proposal['sr_no']}", GREEN
    if change == "Amendment":
        return f"Covered in {proposal['target_test_code']} — amended", AMBER
    if change == "Deletion":
        target = proposal["target_test_code"] or "affected tests"
        return f"Deletion — {target} withdrawn", RED
    return "Information", GREY


def build(document_id: int) -> Path:
    doc_row = store.one("SELECT * FROM documents WHERE id = ?", (document_id,))
    clauses = store.query(
        "SELECT * FROM clauses WHERE document_id = ? ORDER BY sequence", (document_id,))
    proposals = {p["clause_id"]: p for p in store.query(
        "SELECT * FROM proposals WHERE document_id = ?", (document_id,))}

    document = Document()
    for section in document.sections:
        section.left_margin = section.right_margin = Inches(0.7)

    title = document.add_paragraph()
    run = title.add_run("Audit Checklist Working Document")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    run = subtitle.add_run(doc_row["title"] or doc_row["filename"])
    run.font.size = Pt(12)
    run.font.color.rgb = GREY

    meta = document.add_paragraph()
    run = meta.add_run(
        f"Source: {doc_row['source']}   ·   File: {doc_row['filename']}"
        f"{'   ·   Dated: ' + doc_row['doc_date'] if doc_row['doc_date'] else ''}")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY

    document.add_paragraph()
    heading = document.add_paragraph()
    run = heading.add_run("Circular text and audit working notes")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(4.9)
    table.columns[1].width = Inches(2.2)

    header = table.rows[0].cells
    for cell, label in zip(header, ("Clause", "Audit working note")):
        para = cell.paragraphs[0]
        run = para.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY

    for clause in clauses:
        proposal = proposals.get(clause["id"])
        note, colour = _note(proposal)

        row = table.add_row().cells
        left = row[0].paragraphs[0]
        ref = left.add_run(f"{clause['clause_ref']}  ")
        ref.bold = True
        ref.font.size = Pt(9)
        ref.font.color.rgb = NAVY
        body = left.add_run(clause["text"])
        body.font.size = Pt(9.5)

        right = row[1].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = right.add_run(note)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = colour
        if proposal and proposal["change_type"] != "No action":
            detail = right.add_run("\n" + (proposal["proposed_test_description"]
                                           or proposal["rationale"])[:230])
            detail.font.size = Pt(8)
            detail.font.color.rgb = GREY

    config.ensure_dirs()
    # Name the file after the SOURCE FILE, not the detected title — several documents
    # can share a title line and would silently overwrite each other.
    stem = Path(doc_row["filename"]).stem
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in stem)[:55].strip()
    path = config.OUTPUT_DIR / f"{doc_row['id']:02d}_{safe or 'document'}_working.docx"
    document.save(path)
    return path


def build_all() -> list[Path]:
    # Duplicates were never processed, and a document that could not be read has no
    # content to put in a working file — its error shows in the queue instead.
    rows = store.query(
        "SELECT id FROM documents WHERE status = 'ingested'")
    return [build(r["id"]) for r in rows]
